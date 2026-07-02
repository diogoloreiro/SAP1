#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Monta o RELATORIO FINAL unico do SAP-1:
  Capa + Fundamentacao teorica  (front matter, aqui)
  Parte I  - Projeto e implementacao   (corpo de gerar_relatorio_implementacao.py)
  Parte II - Verificacao por simulacao  (corpo de gerar_relatorio_simulacao.py)
Os corpos sao gerados sem capa (SAP1_NOCAPA=1) e compostos com docxcompose.
Saida: Relatorio_Final_SAP1.docx
"""
import os, subprocess
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer

HERE = os.path.dirname(os.path.abspath(__file__))
IMG  = os.path.abspath(os.path.join(HERE, "..", "apresentacao", "img"))
OUT  = os.path.join(HERE, "Relatorio_Final_SAP1.docx")
IMPL_BODY = os.path.join(HERE, "_impl_body.docx")
SIM_BODY  = os.path.join(HERE, "_sim_body.docx")

DARK = RGBColor(0x0d, 0x3b, 0x66); GREY = RGBColor(0x55, 0x55, 0x55)

# ---------------------------------------------------------------
# 1) gerar os corpos (sem capa)
# ---------------------------------------------------------------
env = dict(os.environ, SAP1_NOCAPA="1")
for script in ("gerar_relatorio_implementacao.py", "gerar_relatorio_simulacao.py"):
    subprocess.run(["python3", os.path.join(HERE, script)], env=env, check=True,
                   cwd=HERE, stdout=subprocess.DEVNULL)
print("corpos gerados.")

# ---------------------------------------------------------------
# 2) front matter (capa + fundamentacao teorica)
# ---------------------------------------------------------------
doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.15
for s, sz in (("Heading 1", 16), ("Heading 2", 13)):
    st = doc.styles[s]; st.font.color.rgb = DARK; st.font.size = Pt(sz); st.font.name = "Calibri"

def p(text="", bold=False, italic=False, size=None, color=None, align=None):
    par = doc.add_paragraph()
    if text:
        r = par.add_run(text); r.bold = bold; r.italic = italic
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = color
    if align is not None: par.alignment = align
    return par

def bullet(text):
    par = doc.add_paragraph(style="List Bullet")
    for i, ch in enumerate(text.split("**")):
        r = par.add_run(ch); r.bold = (i % 2 == 1)
    return par

def figure(path, caption, width_cm=15.5):
    if not os.path.exists(path): return
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(path, width=Cm(width_cm))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(10)

def table(headers, rows, widths=None, fs=10):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ""
        rr = c.paragraphs[0].add_run(h); rr.bold = True; rr.font.size = Pt(fs)
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ""; rr = cells[j].paragraphs[0].add_run(str(v)); rr.font.size = Pt(fs)
    if widths:
        for j, w in enumerate(widths):
            for r_ in t.rows: r_.cells[j].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ---- CAPA ----
for _ in range(3): doc.add_paragraph()
p("Processador SAP-1", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=32, color=DARK)
p("Relatório Final", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, color=GREY)
doc.add_paragraph()
p("Projeto e implementação em Verilog · Verificação por simulação",
  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12)
p("Plataforma-alvo: Terasic DE10-Lite (Intel MAX 10 10M50DAF484C7G) · Quartus Prime Lite 20.1",
  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10.5, color=GREY)
for _ in range(6): doc.add_paragraph()
p("Disciplina: __________________________", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
p("Professor(a): __________________________", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
p("Integrantes:", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
p("Diogo Loreiro Campos", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
p("Ivan Marcos", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
p("Marjorie Lobo", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_paragraph()
p("Data: ____ / ____ / ______", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_page_break()

# ---- SUMÁRIO (nota) ----
doc.add_heading("Sumário", level=1)
p("Para inserir o índice automático no Word: posicione o cursor aqui e use "
  "Referências ▸ Sumário ▸ Sumário Automático. Os títulos deste documento usam os "
  "estilos Título 1/2, então o índice é gerado e numerado automaticamente.", italic=True,
  color=GREY)
doc.add_page_break()

# ---- RESUMO ----
doc.add_heading("Resumo", level=1)
p("Este trabalho apresenta o projeto, a implementação em Verilog e a verificação do "
  "processador didático SAP-1 (Simple-As-Possible), com execução na placa FPGA Terasic "
  "DE10-Lite. O processador — de 8 bits, com barramento único, memória de 16 palavras e "
  "cinco instruções — foi descrito de forma modular (um arquivo por bloco), tendo a "
  "unidade de controle modelada como máquina de estados que gera uma palavra de controle "
  "de 12 bits. A verificação foi feita por simulação no ModelSim, com 14 testbenches "
  "auto-verificáveis (por componente e do sistema completo) e análise de formas de onda. "
  "Quatro programas de teste — potência, uma expressão aritmética, multiplicação e "
  "divisão — produziram os resultados esperados (27, 18, 12 e 3) e o processador parou "
  "corretamente em todos. O resultado é um núcleo sintetizável, determinístico na "
  "partida e documentado, que roda de verdade na placa.")
p("Palavras-chave: SAP-1; Verilog; FPGA; DE10-Lite; máquina de estados; simulação.",
  italic=True, color=GREY)

# ---- OBJETIVOS ----
doc.add_heading("Objetivos", level=1)
p("Objetivo geral:", bold=True)
p("Implementar e validar o processador SAP-1 em Verilog, do projeto lógico até a "
  "execução na FPGA DE10-Lite, compreendendo na prática o funcionamento interno de um "
  "processador.")
p("Objetivos específicos:", bold=True)
bullet("Compreender a arquitetura do SAP-1 (datapath, barramento e unidade de controle).")
bullet("Descrever cada bloco como um módulo Verilog, com boas práticas de projeto para FPGA.")
bullet("Modelar a unidade de controle como uma máquina de estados que gera a palavra de controle.")
bullet("Verificar o comportamento por simulação — componente a componente e o sistema completo.")
bullet("Executar programas reais na placa DE10-Lite e observar o funcionamento nos displays.")
bullet("Documentar o projeto, as decisões de implementação e os resultados.")

# ---- METODOLOGIA ----
doc.add_heading("Metodologia", level=1)
p("O trabalho seguiu uma abordagem incremental e verificável, em quatro etapas:")
bullet("**Projeto modular:** cada bloco do datapath e o controle viraram arquivos "
       "Verilog separados, integrados por um módulo de topo através do barramento.")
bullet("**Controle por máquina de estados:** a sequência de microoperações foi "
       "centralizada em um contador de anel (T1–T6) que emite a palavra de controle.")
bullet("**Verificação por simulação:** antes de ir à placa, cada módulo foi exercitado "
       "por um testbench auto-verificável e o sistema completo foi testado com programas "
       "reais, analisando as formas de onda no ModelSim.")
bullet("**Execução na FPGA:** o núcleo, sem alterações, foi adaptado à DE10-Lite (divisor "
       "de clock, botão com antirruído, displays) e gravado para execução real.")
p("Ferramentas utilizadas: Quartus Prime Lite 20.1 (síntese e gravação), ModelSim ASE "
  "(simulação), Icarus Verilog (referência cruzada) e Python/matplotlib (diagramas). A "
  "descrição detalhada da verificação está na Parte II.")

doc.add_page_break()

# ---- FUNDAMENTAÇÃO TEÓRICA ----
doc.add_heading("Fundamentação teórica", level=1)
p("Esta seção reúne os conceitos que embasam o projeto, antes das partes de "
  "implementação e verificação. O objetivo é tornar o relatório autocontido: quem não "
  "conhece o SAP-1 encontra aqui o vocabulário necessário para acompanhar o restante.")

doc.add_heading("O processador SAP-1", level=2)
p("O SAP-1 (Simple-As-Possible) é um processador didático proposto por Albert Malvino "
  "em Digital Computer Electronics. Sua finalidade não é desempenho, mas ensino: é o "
  "menor computador que ainda ilustra todos os princípios de um processador real. "
  "Apesar de ter só 8 bits, um único barramento e cinco instruções, ele executa o "
  "mesmo ciclo fundamental de qualquer CPU — buscar uma instrução na memória, "
  "decodificá-la e executá-la — repetidamente, até parar.")
figure(os.path.join(IMG, "datapath.png"),
       "Visão geral do datapath do SAP-1 (detalhado na Parte I).", 15.5)

doc.add_heading("Arquitetura de Von Neumann", level=2)
p("O SAP-1 adota a arquitetura de Von Neumann: instruções e dados residem na mesma "
  "memória e trafegam pelo mesmo barramento. A principal consequência prática, "
  "explorada ao longo do relatório, é que os 16 endereços da memória são compartilhados "
  "entre o programa e os dados — um programa maior deixa menos espaço para dados. Essa "
  "unificação simplifica o hardware (uma só memória, um só barramento) ao custo de "
  "limitar o tamanho dos programas.")

doc.add_heading("Ciclo de instrução: busca, decodificação e execução", level=2)
p("Toda instrução passa por três fases conceituais. Na busca, o endereço da instrução "
  "(guardado no contador de programa) é enviado à memória e a instrução lida é levada "
  "ao registrador de instrução. Na decodificação, o campo de operação (opcode) é "
  "interpretado pela unidade de controle. Na execução, a unidade de controle gera a "
  "sequência de sinais que realiza a operação (carregar um dado, somar, subtrair, "
  "exibir, parar). No SAP-1 essas fases são realizadas em seis passos de tempo (T1–T6): "
  "os três primeiros são a busca, iguais para toda instrução, e os três últimos a "
  "execução, que depende do opcode.")

doc.add_heading("Conjunto de instruções (ISA)", level=2)
p("O SAP-1 possui cinco instruções. Cada palavra de 8 bits divide-se em dois campos de "
  "4 bits: o opcode (o que fazer) e o operando (um endereço de memória). O endereçamento "
  "é direto — o operando indica onde está o dado, não o próprio valor.")
table(
    ["Instrução", "Opcode", "Significado"],
    [
        ["LDA end", "0000", "carrega o acumulador com o conteúdo de 'end'"],
        ["ADD end", "0001", "soma ao acumulador o conteúdo de 'end'"],
        ["SUB end", "0010", "subtrai do acumulador o conteúdo de 'end'"],
        ["OUT", "1110", "copia o acumulador para o registrador de saída"],
        ["HLT", "1111", "para o processador"],
    ],
    widths=[3.0, 2.2, 9.8], fs=10,
)
p("Não há instruções de escrita em memória, de desvio ou de laço. Isso significa que "
  "os programas são sequências lineares e que operações como multiplicação e divisão "
  "precisam ser expressas por somas e subtrações repetidas — tema retomado na Parte II.")

doc.add_heading("Unidade de controle e microoperações", level=2)
p("A execução de cada instrução se decompõe em microoperações — pequenas transferências "
  "entre registradores comandadas por sinais de controle. A unidade de controle do "
  "SAP-1 é do tipo cabeada (hardwired): a cada passo de tempo ela emite uma palavra de "
  "controle cujos bits habilitam quem escreve e quem lê no barramento. Como essa palavra "
  "depende apenas do passo de tempo e do opcode — nunca do dado —, o comportamento é "
  "determinístico. A construção dessa unidade como máquina de estados é o núcleo da "
  "Parte I.")

doc.add_heading("Representação numérica", level=2)
p("Todos os valores são de 8 bits sem sinal explícito (0 a 255). A subtração usa "
  "complemento de dois, e tanto somas quanto subtrações \"dão a volta\" em 256 "
  "(aritmética modular) — por exemplo, 200 + 100 resulta em 44. Esse comportamento é "
  "inerente à largura fixa de 8 bits e é verificado em simulação na Parte II.")

# divisória Parte I
doc.add_page_break()
doc.add_heading("Parte I — Projeto e implementação", level=1)
p("Esta parte descreve a estrutura do projeto, as decisões de implementação, a máquina "
  "de estados, a palavra de controle e o código-fonte comentado. A numeração de seções "
  "a seguir é interna a esta parte.", italic=True, color=GREY)

doc.save(OUT)

# ---------------------------------------------------------------
# 3) compor: front + Parte I (impl) + Parte II (sim)
# ---------------------------------------------------------------
master = Document(OUT)
composer = Composer(master)
composer.append(Document(IMPL_BODY))

# divisória Parte II
composer.doc.add_page_break()
h = composer.doc.add_heading("Parte II — Verificação por simulação", level=1)
pp = composer.doc.add_paragraph()
r = pp.add_run("Esta parte descreve a metodologia de simulação, a verificação de cada "
               "componente por formas de onda e a execução dos quatro programas de teste. "
               "A numeração de seções a seguir é interna a esta parte.")
r.italic = True; r.font.color.rgb = GREY

composer.append(Document(SIM_BODY))

# ---------------------------------------------------------------
# 4) Execução na placa (DE10-Lite) — espaço reservado p/ fotos
# ---------------------------------------------------------------
m = composer.doc
m.add_page_break()
m.add_heading("Execução na placa (DE10-Lite)", level=1)
pp = m.add_paragraph(
    "Além da simulação, os programas foram gravados e executados na FPGA DE10-Lite. "
    "Nos displays de 7 segmentos é possível acompanhar, em tempo real, o estado atual "
    "(T1–T6), a instrução (como letra), o acumulador e o resultado; o LED de HLT acende "
    "quando o processador para. As fotos abaixo registram cada programa em execução.")
note = m.add_paragraph()
rn = note.add_run("Instrução: substitua cada quadro cinza por uma foto do respectivo "
                  "programa rodando na placa (clique no quadro e insira a imagem).")
rn.italic = True; rn.font.size = Pt(9); rn.font.color.rgb = GREY

def _foto_cell(cell, titulo, legenda):
    # sombreamento cinza
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "ECEFF1")
    tcPr.append(shd)
    cell.text = ""
    p1 = cell.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("[ espaço reservado para foto ]")
    r1.font.size = Pt(10); r1.font.color.rgb = RGBColor(0x90, 0x9a, 0xa0)
    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(titulo); r2.bold = True; r2.font.size = Pt(10)
    p3 = cell.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(legenda); r3.font.size = Pt(9); r3.font.color.rgb = GREY

progs = [
    ("Programa 1 — 3³", "resultado esperado no visor: 1B (27)"),
    ("Programa 2 — expressão", "resultado esperado no visor: 12 (18)"),
    ("Programa 3 — 3 × 4", "resultado esperado no visor: 0C (12)"),
    ("Programa 4 — 12 ÷ 4", "resultado esperado no visor: 03 (3)"),
]
tbl = m.add_table(rows=2, cols=2); tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (tit, leg) in enumerate(progs):
    cell = tbl.rows[i // 2].cells[i % 2]
    _foto_cell(cell, tit, leg)
    # altura minima da linha (~5,5 cm) para caber a foto
    tr = tbl.rows[i // 2]._tr
    trPr = tr.get_or_add_trPr()
    trh = OxmlElement("w:trHeight"); trh.set(qn("w:val"), "3100"); trh.set(qn("w:hRule"), "atLeast")
    trPr.append(trh)
for col in tbl.columns:
    for c in col.cells:
        c.width = Cm(7.6)

composer.save(OUT)
print("gerado:", OUT)

# limpa corpos temporarios
for f in (IMPL_BODY, SIM_BODY):
    try: os.remove(f)
    except OSError: pass
