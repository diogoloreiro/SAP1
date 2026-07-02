#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Relatorio de IMPLEMENTACAO (RTL) do SAP-1 em DOCX - versao aprofundada.
Estrutura, decisoes de projeto no nivel de hardware, disciplina de temporizacao,
maquina de estados (duas implementacoes), tabela completa da palavra de controle,
rastreamento ciclo-a-ciclo, codigo-fonte completo e camada de placa (fpga/).
Requer: python-docx. Imagens em ../fsm/ e ../apresentacao/img/.
Saida: Relatorio_Implementacao_SAP1.docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
IMG  = os.path.abspath(os.path.join(HERE, "..", "apresentacao", "img"))
FSM  = os.path.abspath(os.path.join(HERE, "..", "fsm"))
OUT  = os.path.join(HERE, "Relatorio_Implementacao_SAP1.docx")
NOCAPA = os.environ.get("SAP1_NOCAPA") == "1"   # gera corpo sem capa (p/ compor)
if NOCAPA:
    OUT = os.path.join(HERE, "_impl_body.docx")

DARK  = RGBColor(0x0d, 0x3b, 0x66)
GREY  = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1b, 0x5e, 0x20)

# ---- cache de imagens reduzidas (abre mais rapido no Word/LibreOffice) ----
import tempfile, hashlib
from PIL import Image
CACHE = tempfile.mkdtemp(prefix="sap1img_")
def _prep(path, maxw=1200):
    if not os.path.exists(path):
        return path
    im = Image.open(path).convert("RGB")
    if im.width <= maxw:
        return path
    h = int(im.height * maxw / im.width)
    out = os.path.join(CACHE, hashlib.md5(path.encode()).hexdigest() + ".png")
    im.resize((maxw, h), Image.LANCZOS).save(out, optimize=True, compress_level=9)
    return out

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.13
for s, sz in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
    st = doc.styles[s]; st.font.color.rgb = DARK; st.font.size = Pt(sz); st.font.name = "Calibri"

def p(text="", bold=False, italic=False, size=None, color=None, align=None):
    par = doc.add_paragraph()
    if text:
        r = par.add_run(text); r.bold = bold; r.italic = italic
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = color
    if align is not None: par.alignment = align
    return par

def rich(parts):
    """parts: lista de (texto, bold). Um paragrafo com formatacao mista."""
    par = doc.add_paragraph()
    for txt, b in parts:
        r = par.add_run(txt); r.bold = b
    return par

def bullet(text, sub=False):
    par = doc.add_paragraph(style="List Bullet 2" if sub else "List Bullet")
    for i, chunk in enumerate(text.split("**")):
        r = par.add_run(chunk); r.bold = (i % 2 == 1)
    return par

def lead(label, text, color=None):
    par = doc.add_paragraph()
    r = par.add_run(label + " "); r.bold = True
    if color: r.font.color.rgb = color
    if text: par.add_run(text)
    return par

def code(lines, caption=None, fs=8.3):
    if caption:
        cp = doc.add_paragraph()
        r = cp.add_run(caption); r.bold = True; r.font.size = Pt(9)
        r.font.name = "Consolas"; r.font.color.rgb = DARK
        cp.paragraph_format.space_after = Pt(0)
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(0.35)
    par.paragraph_format.space_before = Pt(2); par.paragraph_format.space_after = Pt(8)
    r = par.add_run("\n".join(lines))
    r.font.name = "Consolas"; r.font.size = Pt(fs)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F4F7")
    par._p.get_or_add_pPr().append(shd)
    return par

FIGN = [0]
def figure(path, caption, width_cm=15.5):
    if not os.path.exists(path):
        p(f"[figura ausente: {os.path.basename(path)}]", italic=True, color=GREY); return
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(_prep(path), width=Cm(width_cm))
    FIGN[0] += 1
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figura {FIGN[0]} — {caption}")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(10)

def table(headers, rows, widths=None, fs=10, mono_cols=()):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False  # layout fixo (tblLayout=fixed): evita recalculo pesado ao abrir
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ""
        rr = c.paragraphs[0].add_run(h); rr.bold = True; rr.font.size = Pt(fs)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            rr = cells[j].paragraphs[0].add_run(str(val)); rr.font.size = Pt(fs)
            if j in mono_cols: rr.font.name = "Consolas"; rr.font.size = Pt(fs-0.5)
    if widths:
        for j, w in enumerate(widths):
            for r_ in t.rows: r_.cells[j].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# =====================================================================
# CAPA
# =====================================================================
if not NOCAPA:
    for _ in range(3): doc.add_paragraph()
    p("Processador SAP-1", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=30, color=DARK)
    p("Relatório de Implementação em Verilog (RTL)",
      align=WD_ALIGN_PARAGRAPH.CENTER, size=16, color=GREY)
    doc.add_paragraph()
    p("Arquitetura · Decisões de projeto · Temporização · Máquina de estados · Código-fonte",
      align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11.5)
    p("Plataforma-alvo: Terasic DE10-Lite (Intel MAX 10 10M50DAF484C7G) · Quartus Prime Lite 20.1",
      align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10.5, color=GREY)
    for _ in range(6): doc.add_paragraph()
    p("Disciplina: __________________________", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    p("Integrantes: __________________________", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    p("Data: ____ / ____ / ______", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    doc.add_page_break()

# =====================================================================
# 1. INTRODUÇÃO
# =====================================================================
doc.add_heading("1. Introdução", level=1)
p("Este relatório documenta a implementação em Verilog do processador SAP-1 "
  "(Simple-As-Possible): como o projeto está organizado, quais decisões de projeto "
  "foram tomadas — e por quê, no nível de hardware —, como a temporização foi disciplinada "
  "em duas bordas de clock, como a unidade de controle foi modelada por máquina de "
  "estados, e o código-fonte comentado. Serve de referência técnica do que foi "
  "construído, complementando o relatório de verificação por simulação.")
p("Resumo das especificações implementadas:")
table(
    ["Parâmetro", "Valor", "Onde"],
    [
        ["Largura de dados", "8 bits", "barramento W, A, B, ULA, saída"],
        ["Largura de endereço", "4 bits", "PC e MAR (16 posições)"],
        ["Memória", "16 × 8 bits (ROM em execução)", "ram_16x8.v"],
        ["Registradores de dados", "2 (acumulador A e B)", "accumulator.v, register_b.v"],
        ["ULA", "soma e subtração, 8 bits", "adder_subtractor.v"],
        ["Instruções", "5 (LDA, ADD, SUB, OUT, HLT)", "decodificadas no controle"],
        ["Estados de controle", "6 (T1–T6) + IDLE de partida", "controller_sequencer.v"],
        ["Palavra de controle", "12 bits (CON)", "gerada a cada estado"],
    ],
    widths=[3.8, 5.4, 6.4], fs=9.5,
)

# =====================================================================
# 2. ESTRUTURA DO PROJETO
# =====================================================================
doc.add_heading("2. Estrutura do projeto", level=1)
p("O projeto separa o núcleo sintetizável (rtl) da camada específica de placa (fpga), "
  "da verificação (tb_model) e da documentação:")
code([
    "SAP1/",
    "|- rtl/         nucleo do processador (11 modulos, sinteticaveis, sem I/O de placa)",
    "|- fpga/        camada da placa DE10-Lite (top de placa, clock, debounce, displays)",
    "|- tb_model/    14 testbenches auto-verificaveis + scripts de onda (ModelSim)",
    "|- programas/   programas de teste em .txt (1..4) para colar na RAM",
    "|- fsm/         diagramas da maquina de estados (Python/matplotlib)",
    "|- docs/        documentacao e relatorios (este arquivo)",
    "|- apresentacao/ slides e imagens (datapath, ondas)",
    "|- sap1_top.qsf, SAP1.qpf   projeto Quartus (pinos, arquivos)",
])
rich([("A separação ", False), ("rtl × fpga", True),
      (" é deliberada e tem efeito prático: o núcleo em ", False), ("rtl", True),
      (" não conhece displays, botões nem o clock de 50 MHz da placa — suas únicas "
       "entradas são ", False), ("clk", True), (" e ", False), ("clr_bar", True),
      (". Quem adapta o núcleo ao hardware real (divide o clock, filtra o botão, "
       "aciona os sete segmentos) é a pasta ", False), ("fpga", True),
      (". Assim, o mesmo núcleo é simulado e sintetizado sem alteração.", False)])

doc.add_heading("2.1. Módulos do núcleo (rtl)", level=2)
table(
    ["Módulo", "Papel", "Natureza"],
    [
        ["sap1_top.v", "Topo: barramento W (mux) + instâncias de todos os blocos", "estrutural"],
        ["program_counter.v", "PC — contador de 4 bits do endereço de instrução", "sequencial"],
        ["mar.v", "Registrador de endereço apresentado à memória", "sequencial"],
        ["ram_16x8.v", "Memória 16×8, leitura combinacional (ROM em execução)", "combinacional"],
        ["instruction_register.v", "IR — guarda a instrução e a fatia em opcode|operando", "sequencial"],
        ["accumulator.v", "Acumulador A (registrador de trabalho)", "sequencial"],
        ["register_b.v", "Registrador B (segundo operando da ULA)", "sequencial"],
        ["adder_subtractor.v", "ULA — soma/subtração em 8 bits (complemento de dois)", "combinacional"],
        ["output_register.v", "Registrador de saída (visor)", "sequencial"],
        ["controller_sequencer.v", "Controle usado: anel one-hot + palavra de controle", "sequencial"],
        ["controller_fsm.v", "Controle alternativo: FSM clássica de 3 processos", "sequencial"],
    ],
    widths=[4.4, 8.6, 2.4], fs=9.3,
)

doc.add_heading("2.2. Módulos da camada de placa (fpga)", level=2)
table(
    ["Módulo", "Papel"],
    [
        ["sap1_fpga.v", "Top de placa: junta núcleo, clock, reset, botões e displays"],
        ["clock_divider.v", "Divide 50 MHz para ~1 Hz (execução visível)"],
        ["debouncer.v", "Antirruído do botão (clock manual passo a passo)"],
        ["seg7.v", "Decodificador hexadecimal → 7 segmentos"],
        ["seg7_instr.v", "Decodificador opcode → letra (L/A/5/o/H)"],
    ],
    widths=[3.6, 11.8], fs=9.5,
)
figure(os.path.join(IMG, "module_hierarchy.png"),
       "Hierarquia: o top de placa instancia a camada fpga e o núcleo sap1_top, que "
       "por sua vez instancia os 9 blocos do datapath e controle.", 16.0)
p("A figura evidencia a separação de responsabilidades: o núcleo sap1_top (azul) apenas "
  "instancia os nove blocos do datapath e do controle, sem conhecer nada da placa; toda "
  "a adaptação ao hardware (clock, botão, displays) fica na camada fpga (laranja), sob o "
  "top de placa. Note ainda que controller_fsm.v é uma alternativa ao "
  "controller_sequencer.v — mesmas portas, trocável no topo.")

# =====================================================================
# 3. ARQUITETURA E CONVENÇÕES
# =====================================================================
doc.add_heading("3. Arquitetura e convenções de sinais", level=1)
figure(os.path.join(IMG, "datapath.png"),
       "Datapath do SAP-1: barramento W, blocos de dados e unidade de controle.", 16.0)
p("Todos os blocos se comunicam por um único barramento de 8 bits (barramento W). A "
  "regra de uso do barramento é rígida e é o que garante a coerência do datapath:")
bullet("**Em cada estado, no máximo um bloco escreve** no barramento (um único "
       "\"enable\" de saída ativo); um ou mais blocos podem **ler** simultaneamente.")
bullet("Fontes de 4 bits (PC e operando do IR) são **estendidas com zeros** à esquerda "
       "para 8 bits: {4'b0000, pc_out}. O MAR, por sua vez, lê apenas os 4 bits baixos "
       "do barramento (w_bus[3:0]).")

doc.add_heading("3.1. Convenção de nomes e níveis ativos", level=2)
p("A nomenclatura segue a tradição do SAP-1 e é seguida à risca no código, o que evita "
  "erros de polaridade:")
bullet("Sinais **ativos em alto** (1 = ativo): Cp, Ep, Ea, Su, Eu. São \"habilitações\" "
       "de contagem/saída no barramento.")
bullet("Sinais **ativos em baixo** (0 = ativo), grafados com til/sufixo _bar: ~Lm, ~CE, "
       "~Li, ~Ei, ~La, ~Lb, ~Lo. São as cargas de registradores e a habilitação da RAM.")
rich([("Consequência importante: o ", False), ("repouso não é tudo-zero", True),
      (". Como as cargas são ativas em baixo, a palavra de repouso precisa manter esses "
       "bits em 1. Por isso a palavra ociosa é ", False),
      ("12'b0011_1110_0011", True), (", e não zero.", False)])

doc.add_heading("3.2. Mapa geral das entidades", level=2)
p("Antes de detalhar a temporização e o código, esta é a visão de conjunto: o que cada "
  "bloco faz, o que ele recebe e o que entrega. Todos os registradores compartilham o "
  "mesmo padrão (carga ativa em baixo, reset assíncrono), variando apenas a largura e o "
  "nome do sinal de carga.")
figure(os.path.join(IMG, "entity_map.png"),
       "Mapa geral das entidades: função e sinais de comando de cada bloco.", 15.5)
bullet("**Barramento W (8 bits)** — a via única por onde os dados circulam. Em cada "
       "passo, um bloco escreve nele (habilitado por um sinal de saída) e um ou mais "
       "leem. É implementado por multiplexador (não tri-state).")
bullet("**Contador de Programa (PC)** — guarda o endereço da próxima instrução (4 bits). "
       "Recebe: Cp (incrementa) e Ep (coloca o endereço no barramento). Entrega: o "
       "endereço corrente. É quem faz o programa avançar.")
bullet("**Registrador de Endereço (MAR)** — retém o endereço que a memória vai usar. "
       "Recebe do barramento (4 bits) quando ~Lm ativa; entrega o endereço à RAM. "
       "Carregado duas vezes por instrução (no T1 e no T4).")
bullet("**Memória RAM 16×8** — guarda programa e dados (16 palavras de 8 bits). "
       "Endereçada pelo MAR; entrega o conteúdo no barramento quando ~CE ativa. Somente "
       "leitura em execução (ROM inicializada).")
bullet("**Registrador de Instrução (IR)** — captura a instrução vinda da memória (~Li) "
       "e a separa em opcode (bits altos, para o controle) e operando (bits baixos, que "
       "vai ao barramento quando ~Ei ativa).")
bullet("**Acumulador (A)** — registrador de trabalho de 8 bits: acumula os resultados. "
       "Recebe do barramento (~La) e o entrega tanto continuamente à ULA quanto ao "
       "barramento (Ea). É o registrador mais movimentado do processador.")
bullet("**Registrador B** — guarda o segundo operando da ULA (~Lb). Diferente de A, "
       "serve apenas de \"copo\" temporário durante somas e subtrações.")
bullet("**ULA (somador/subtrator)** — bloco combinacional: calcula A+B ou A−B (conforme "
       "Su) e coloca o resultado no barramento quando Eu ativa. Não guarda estado.")
bullet("**Registrador de Saída** — recebe o acumulador (~Lo) e o mantém estável para o "
       "visor (LEDs/displays), sem afetar o restante do datapath.")
bullet("**Controlador/Sequenciador** — o \"maestro\": a máquina de estados que percorre "
       "T1–T6 e, a cada passo, emite a palavra de controle de 12 bits que aciona todos "
       "os sinais acima. Recebe o opcode do IR e decide a sequência de microoperações.")

# =====================================================================
# 4. DISCIPLINA DE TEMPORIZAÇÃO (DUAS BORDAS)
# =====================================================================
doc.add_heading("4. Disciplina de temporização em duas bordas", level=1)
p("A decisão de temporização mais estrutural do projeto é usar as duas bordas do clock, "
  "com papéis distintos:")
bullet("O **contador de estados avança na borda de descida** (negedge) do clock.")
bullet("Os **registradores do datapath carregam na borda de subida** (posedge).")
p("O motivo é eliminar uma condição de corrida entre \"decidir o que fazer\" e "
  "\"fazer\". Ao trocar de estado no negedge, a nova palavra de controle já está estável "
  "e propagada quando chega o posedge seguinte — instante em que os registradores "
  "amostram o barramento. Se estado e dados mudassem na mesma borda, o registrador "
  "poderia amostrar valores da microoperação errada. Na prática, cada estado Tn ocupa "
  "meio período \"preparando\" o barramento e o posedge no meio dele efetiva a "
  "transferência. Essa relação é o que torna previsível a leitura das ondas (o "
  "acumulador de um ADD só assenta no T6 e aparece assentado já no T1 seguinte).")
figure(os.path.join(IMG, "timing_two_edge.png"),
       "Duas bordas: o estado avança no negedge (bordas) e os registradores carregam no "
       "posedge (meio de cada estado).", 16.0)
p("A figura resume essa disciplina para o ciclo de busca. Nas linhas azuis "
  "(bordas de descida) o estado muda — T1, T2, T3 —; nas linhas laranja (bordas de "
  "subida, no meio de cada estado) os registradores amostram o barramento: MAR ← PC em "
  "T1, PC ← PC+1 em T2 e IR ← RAM em T3. Como a palavra de controle de cada estado já "
  "está estável antes do posedge correspondente, não há ambiguidade sobre qual "
  "transferência ocorre.")

# =====================================================================
# 5. DECISÕES DE IMPLEMENTAÇÃO (nível de hardware)
# =====================================================================
doc.add_heading("5. Decisões de implementação", level=1)
p("As decisões a seguir são ilustradas com trechos de código. As construções de Verilog "
  "empregadas (assign, always, reg/wire, reset assíncrono, non-blocking, etc.) estão "
  "resumidas na seção 10.1 — vale consultá-la antes se o Verilog não for familiar.",
  italic=True)

doc.add_heading("5.1. Barramento W por multiplexador com prioridade", level=2)
p("Nos livros o barramento costuma ser desenhado com buffers tri-state. Optou-se por um "
  "multiplexador com cadeia de prioridade (if-else aninhado via operador ternário), pela "
  "razão concreta de que FPGAs da família MAX 10 não têm barramentos tri-state internos: "
  "a forma sintetizável de \"vários blocos, um fio\" é justamente o mux. A cadeia testa "
  "os enables em ordem; funcionalmente a ordem é irrelevante porque a unidade de "
  "controle garante exclusão mútua (só um enable ativo por estado). O ramo final "
  "(0000_0000) evita propagação de X na simulação quando, por construção, nenhum enable "
  "está ativo (estados de repouso).")
code([
    "assign w_bus = ep             ? {4'b0000, pc_out}      :  // Ep : PC (4b) -> 8b",
    "               (ce_bar==1'b0) ? ram_data               :  // ~CE: RAM (8b)",
    "               (ei_bar==1'b0) ? {4'b0000, ir_operand}  :  // ~Ei: operando (4b) -> 8b",
    "               ea             ? acc_out                :  // Ea : acumulador",
    "               eu             ? alu_out                :  // Eu : saida da ULA",
    "                                8'b0000_0000;            // repouso (nenhum fala)",
], caption="rtl/sap1_top.v — multiplexador do barramento W")

doc.add_heading("5.2. Padrão de registrador: reset assíncrono + carga com prioridade", level=2)
p("Cinco blocos (A, B, MAR, IR e saída) são o mesmo flip-flop com enable. O reset "
  "aparece na lista de sensibilidade (negedge clr_bar), tornando-o assíncrono — ele "
  "vence qualquer coisa. Em seguida, a carga só ocorre quando o sinal está ativo; se não "
  "houver carga, o bloco simplesmente não atribui, e a síntese infere o \"segura o valor\" "
  "de um flip-flop com clock-enable (não gera latch, pois todo caminho está sob a borda "
  "de clock). Padronizar isso deixou o núcleo uniforme — muda só a largura e o nome da "
  "carga.")
code([
    "always @(posedge clk or negedge clr_bar) begin",
    "    if (!clr_bar)      acc_out <= 8'b0000_0000;  // (1) reset assincrono, prioridade",
    "    else if (!la_bar)  acc_out <= bus_in;        // (2) carrega quando ~La = 0",
    "end                                              // (3) senao: mantem (FF com enable)",
], caption="rtl/accumulator.v — padrão de A, B, MAR, IR e saída")

doc.add_heading("5.3. Memória como ROM inicializada (sem STA)", level=2)
p("O SAP-1 básico não tem instrução de escrita (STA); logo a RAM é somente de leitura "
  "durante a execução. Um bloco initial preenche os 16 endereços e a leitura é "
  "combinacional (assign data_out = mem[addr]), sem clock — um \"asynchronous read ROM\". "
  "Na síntese isso vira memória de inicialização (ROM). Como programa e dados dividem os "
  "mesmos 16 slots (Von Neumann), o tamanho do programa limita o espaço de dados. Trocar "
  "de programa = editar este bloco e recompilar; por isso os quatro programas ficam "
  "prontos em programas/*.txt.")
code([
    "reg [7:0] mem [0:15];",
    "integer i;",
    "initial begin",
    "    for (i = 0; i < 16; i = i + 1) mem[i] = 8'h00;  // zera tudo",
    "    mem[0]  = 8'b0000_1011; // LDA 11   <- opcode|operando",
    "    mem[1]  = 8'b0001_1100; // ADD 12",
    "    // ...                                (programa)",
    "    mem[10] = 8'b1111_0000; // HLT",
    "    mem[11] = 8'd7;         //           (dados)",
    "end",
    "assign data_out = mem[addr];   // leitura assincrona (combinacional)",
], caption="rtl/ram_16x8.v")
figure(os.path.join(IMG, "memory_map.png"),
       "Os 16 slots compartilhados: programa (0–10) e dados (11–15) do Programa 2.", 9.5)
p("A figura mostra essa divisão para o Programa 2: as posições 0 a 10 guardam as "
  "instruções e 11 a 15 os dados. Fica evidente o compromisso imposto pelos 4 bits de "
  "endereço — programa e dados disputam os mesmos 16 slots, de modo que um programa "
  "maior deixa menos espaço para dados.")

doc.add_heading("5.4. ULA combinacional em complemento de dois", level=2)
p("A ULA é uma única expressão combinacional. A subtração usa o complemento de dois "
  "nativo do Verilog e o resultado de 8 bits satura por wraparound (módulo 256). Por ser "
  "combinacional, o resultado precisa estabilizar dentro do estado T6 antes do posedge "
  "que o grava em A — o que é folgado no clock lento da placa. É essa largura fixa de 8 "
  "bits que limita o SAP-1 a valores de 0 a 255.")
code([
    "assign result = su ? (acc - breg)   // Su = 1: subtrai (compl. de dois)",
    "                   : (acc + breg);  // Su = 0: soma",
], caption="rtl/adder_subtractor.v")

doc.add_heading("5.5. HLT por clock-enable, não por gating de clock", level=2)
p("Decisão de robustez central. Parar o processador desligando o clock por lógica "
  "combinacional (clock gating) é má prática em FPGA: introduz glitches no sinal de "
  "clock, viola as restrições de tempo (STA) e pode deixar registradores em estados "
  "inconsistentes. A solução adotada é registrar um flag halted (travado no estado T4, "
  "quando o opcode HLT já foi decodificado em T3) e usá-lo como clock-enable do contador "
  "de estados. Congelado o contador, a palavra de controle vira IDLE, nenhuma carga é "
  "ativada e todos os registradores seguram seus valores naturalmente — o processador "
  "\"para\" com o clock intacto e o estado preservado em T4. No topo, o clock do núcleo "
  "é ligado direto (wire clk_cpu = clk;), sem qualquer gating.")
code([
    "reg  halted;",
    "wire run = ~halted;                       // clock-enable do contador de anel",
    "always @(posedge clk or negedge clr_bar) begin",
    "    if (!clr_bar)                              halted <= 1'b0;",
    "    else if ((ring == T4) && (opcode == HLT))  halted <= 1'b1;  // trava em T4",
    "end",
], caption="rtl/controller_sequencer.v — HALT registrado")

doc.add_heading("5.6. Estado ocioso de partida (IDLE) e recuperação de reset", level=2)
p("O contador inclui um estado de partida IDLE (nenhum T ativo, valor one-hot 000000). "
  "Ele absorve a fase do reset. Sem esse estado, o funcionamento correto dependeria de "
  "soltar o reset exatamente com o clock em nível conveniente — um problema clássico de "
  "recovery/removal do reset assíncrono. Com o IDLE, qualquer que seja a fase do clock ao "
  "soltar o reset, o primeiro negedge útil leva a máquina a T1 de forma determinística. "
  "Na camada de placa, esse cuidado é reforçado por um sincronizador de reset (assert "
  "assíncrono, desassert síncrono), descrito na seção 9.")

# =====================================================================
# 6. MÁQUINA DE ESTADOS
# =====================================================================
doc.add_heading("6. Máquina de estados (unidade de controle)", level=1)

doc.add_heading("6.1. Diagrama de estados (anel)", level=2)
figure(os.path.join(FSM, "fsm_diagram.png"),
       "Contador de anel IDLE → T1 → … → T6 → T1; laço roxo em T4 = congelamento por HLT.", 14.5)

doc.add_heading("6.2. Ciclo de execução por instrução", level=2)
figure(os.path.join(FSM, "fsm_exec_diagram.png"),
       "Busca comum (T1–T3, azul) e ramificação da execução (T4–T6) por opcode.", 16.0)
p("A busca é idêntica para toda instrução; a diferença começa em T4, quando o opcode é "
  "decodificado. LDA usa até T5 (A ← RAM); ADD/SUB usam até T6 (A ← A±B); OUT resolve em "
  "T4; HLT congela em T4.")

doc.add_heading("6.3. Codificação one-hot e a rotação do anel", level=2)
p("O estado é mantido em one-hot de 6 bits (exatamente um bit em 1 por vez, salvo o "
  "IDLE que é todo-zero). One-hot foi escolhido por dois motivos concretos: a decodificação "
  "do estado vira um simples teste de um bit (barato e rápido em LUTs de FPGA), e a "
  "transição para o próximo estado é uma rotação de 1 bit — sem somador nem comparador. "
  "A rotação é feita por concatenação:")
code([
    "// ring_r = {b5,b4,b3,b2,b1,b0}   (T1=000001, T2=000010, ..., T6=100000)",
    "ring_r <= {ring_r[4:0], ring_r[5]};  // desloca a esquerda e traz o bit 5 p/ o 0",
    "// T1(000001)->T2(000010)->T3(000100)->T4(001000)->T5(010000)->T6(100000)->T1",
], caption="rotação one-hot (T6 → T1 fecha o anel)")
p("O contador completo trata os quatro casos — reset, congelado (HLT), partida do IDLE e "
  "rotação — nesta ordem de prioridade:")
code([
    "always @(negedge clk or negedge clr_bar) begin",
    "    if (!clr_bar)             ring_r <= IDLE_ST;                 // reset -> IDLE",
    "    else if (!run)            ring_r <= ring_r;                  // HLT: congela",
    "    else if (ring_r==IDLE_ST) ring_r <= T1;                     // 1o negedge -> T1",
    "    else                      ring_r <= {ring_r[4:0], ring_r[5]}; // rotaciona",
    "end",
], caption="rtl/controller_sequencer.v — contador de anel")

doc.add_heading("6.4. Duas implementações equivalentes do controle", level=2)
p("A unidade de controle existe em dois arquivos, funcionalmente equivalentes (mesmas "
  "portas, mesma tabela de saída), como estudo comparativo:")
table(
    ["Aspecto", "controller_sequencer.v (usado)", "controller_fsm.v (alternativo)"],
    [
        ["Estilo", "contador de anel (shift register)", "FSM clássica de 3 processos"],
        ["Codificação do estado", "one-hot (6 bits)", "binária (3 bits, nomes simbólicos)"],
        ["Próximo estado", "rotação de 1 bit", "case explícito S_T1→S_T2→…"],
        ["Vantagem", "decode/rotação triviais", "leitura didática, estados nomeados"],
    ],
    widths=[3.2, 6.1, 6.1], fs=9.3,
)
p("Trocar uma pela outra é só mudar o nome do módulo instanciado no topo (as portas são "
  "idênticas). O trecho de próximo-estado da versão clássica deixa a sequência explícita:")
code([
    "always @(*) begin",
    "    if (!run) next_state = state;          // HLT: congela",
    "    else case (state)",
    "        S_IDLE: next_state = S_T1;",
    "        S_T1:   next_state = S_T2;   S_T2: next_state = S_T3;",
    "        S_T3:   next_state = S_T4;   S_T4: next_state = S_T5;",
    "        S_T5:   next_state = S_T6;   S_T6: next_state = S_T1;  // fecha o anel",
    "        default: next_state = S_IDLE;",
    "    endcase",
    "end",
], caption="rtl/controller_fsm.v — lógica de próximo estado (2º processo)")

doc.add_heading("6.5. Descrição estado a estado (T1–T6)", level=2)
p("A seguir, o papel de cada passo de tempo. Os três primeiros formam a busca e são "
  "idênticos para toda instrução; os três últimos formam a execução e dependem do "
  "opcode. Em todos, a palavra de controle do estado é preparada no negedge e a "
  "transferência se efetiva no posedge seguinte.")

lead("T1 — envia o endereço da instrução (MAR ← PC).",
     "O contador de programa é habilitado a escrever no barramento (Ep) e o MAR a "
     "carregar (~Lm). Ao fim do estado, o MAR contém o endereço da instrução a buscar. "
     "É o passo que \"aponta\" para a memória.")
lead("T2 — incrementa o contador (PC ← PC + 1).",
     "Apenas Cp fica ativo: o PC soma 1, passando a apontar para a próxima instrução. "
     "O barramento não é usado para transferência — este passo apenas adianta o PC "
     "enquanto o endereço atual já está seguro no MAR.")
lead("T3 — busca e decodifica (IR ← RAM).",
     "A RAM é habilitada a escrever no barramento (~CE) e o IR a carregar (~Li). A "
     "instrução apontada pelo MAR é lida e guardada no IR, que imediatamente expõe o "
     "opcode ao controle. A partir daqui a máquina \"sabe\" qual é a instrução — a busca "
     "terminou.")
lead("T4 — primeiro passo da execução (depende do opcode).",
     "Para LDA/ADD/SUB, o operando do IR vai ao MAR (~Ei, ~Lm), preparando a leitura do "
     "dado. Para OUT, o acumulador é copiado ao registrador de saída (Ea, ~Lo) e a "
     "instrução termina aqui. Para HLT, o flag de parada é travado e a máquina congela "
     "em T4.")
lead("T5 — busca o operando (depende do opcode).",
     "Para LDA, o dado lido da RAM vai direto ao acumulador (~CE, ~La) — a instrução "
     "acaba. Para ADD/SUB, o dado vai ao registrador B (~CE, ~Lb), ficando pronto para "
     "a operação aritmética do passo seguinte. Para OUT/HLT, o passo é vazio.")
lead("T6 — conclui a aritmética (só ADD/SUB).",
     "A ULA calcula A+B (ADD) ou A−B (SUB, com Su) e o resultado retorna ao acumulador "
     "(Eu, ~La). Como A só é reescrito neste posedge, o novo valor aparece assentado já "
     "no T1 da instrução seguinte. Para as demais instruções, o passo é vazio.")
p("Depois do T6 (ou do passo em que a instrução termina), o anel volta a T1 e o ciclo "
  "recomeça para a próxima instrução — exceto após HLT, em que a máquina permanece "
  "congelada em T4.")

# =====================================================================
# 7. A PALAVRA DE CONTROLE (COMPLETA)
# =====================================================================
doc.add_heading("7. A palavra de controle de 12 bits", level=1)
p("A cada estado a unidade de controle emite 12 bits (CON) que comandam todo o datapath. "
  "O layout é fixo (MSB → LSB):")
code([
    "CON = { Cp, Ep, ~Lm, ~CE, ~Li, ~Ei, ~La, Ea, Su, Eu, ~Lb, ~Lo }",
    "bit:   11  10   9    8    7    6    5   4   3   2   1    0",
])
figure(os.path.join(IMG, "control_word.png"),
       "Layout dos 12 bits de CON: azul = ativo em alto, laranja = ativo em baixo.", 16.0)
p("Na figura, os bits em azul são ativos em alto (habilitações) e os em laranja, ativos "
  "em baixo (cargas e a habilitação da RAM). Essa distinção visual explica por que a "
  "palavra de repouso mantém 1 nos bits de carga — em vez de zero — e serve de guia para "
  "ler a tabela detalhada a seguir.")
table(
    ["Bit", "Sinal", "Ativo", "Função quando ativo"],
    [
        ["11", "Cp",  "alto",  "incrementa o PC"],
        ["10", "Ep",  "alto",  "PC escreve no barramento"],
        ["9",  "~Lm", "baixo", "carrega o MAR a partir do barramento"],
        ["8",  "~CE", "baixo", "RAM escreve no barramento"],
        ["7",  "~Li", "baixo", "carrega o IR"],
        ["6",  "~Ei", "baixo", "operando do IR no barramento"],
        ["5",  "~La", "baixo", "carrega o acumulador"],
        ["4",  "Ea",  "alto",  "acumulador escreve no barramento"],
        ["3",  "Su",  "alto",  "ULA subtrai (0 = soma)"],
        ["2",  "Eu",  "alto",  "ULA escreve no barramento"],
        ["1",  "~Lb", "baixo", "carrega o registrador B"],
        ["0",  "~Lo", "baixo", "carrega o registrador de saída"],
    ],
    widths=[1.2, 1.8, 2.0, 9.5], fs=9.3,
)
p("A tabela a seguir lista a palavra emitida em cada estado/opcode, os sinais que ficam "
  "ativos e a microoperação resultante. Os valores binários foram transcritos "
  "diretamente do código (e conferidos bit a bit):")
table(
    ["Estado", "Opcode", "CON (12 bits)", "Sinais ativos", "Microoperação"],
    [
        ["IDLE", "—", "0011_1110_0011", "(nenhum)", "repouso"],
        ["T1", "todos", "0101_1110_0011", "Ep, ~Lm", "MAR ← PC"],
        ["T2", "todos", "1011_1110_0011", "Cp", "PC ← PC + 1"],
        ["T3", "todos", "0010_0110_0011", "~CE, ~Li", "IR ← RAM[MAR]"],
        ["T4", "LDA/ADD/SUB", "0001_1010_0011", "~Lm, ~Ei", "MAR ← IR.operando"],
        ["T4", "OUT", "0011_1111_0010", "Ea, ~Lo", "saída ← A"],
        ["T4", "HLT", "(IDLE)", "run ← 0", "congela em T4"],
        ["T5", "LDA", "0010_1100_0011", "~CE, ~La", "A ← RAM[MAR]"],
        ["T5", "ADD/SUB", "0010_1110_0001", "~CE, ~Lb", "B ← RAM[MAR]"],
        ["T6", "ADD", "0011_1100_0111", "~La, Eu", "A ← A + B"],
        ["T6", "SUB", "0011_1100_1111", "~La, Eu, Su", "A ← A − B"],
    ],
    widths=[1.9, 2.4, 3.4, 3.0, 4.0], fs=9.0, mono_cols=(2,),
)
p("Estados não listados para um dado opcode (por exemplo, T5/T6 de OUT) emitem a palavra "
  "IDLE — são \"vazios\", o que explica por que toda instrução ocupa 6 estados mesmo "
  "quando usa menos.", italic=True)

# =====================================================================
# 8. RASTREAMENTO CICLO A CICLO (ADD)
# =====================================================================
doc.add_heading("8. Rastreamento ciclo a ciclo de uma instrução (ADD)", level=1)
p("Para tornar concreta a interação controle × datapath, segue a execução de um ADD "
  "endereço, com o acumulador inicialmente em A e o dado D na posição apontada. Cada "
  "linha é um estado: o controle prepara a palavra (no negedge de entrada do estado) e a "
  "transferência se efetiva no posedge seguinte.")
table(
    ["Estado", "Sinais ativos", "Quem fala no barramento", "Quem carrega (posedge)"],
    [
        ["T1", "Ep, ~Lm", "PC", "MAR ← PC"],
        ["T2", "Cp", "(ninguém)", "PC ← PC + 1"],
        ["T3", "~CE, ~Li", "RAM", "IR ← instrução ADD"],
        ["T4", "~Lm, ~Ei", "operando do IR", "MAR ← endereço do dado"],
        ["T5", "~CE, ~Lb", "RAM", "B ← D (o dado)"],
        ["T6", "~La, Eu", "ULA (A + B)", "A ← A + D"],
    ],
    widths=[1.8, 3.0, 5.0, 4.6], fs=9.5,
)
figure(os.path.join(IMG, "add_flow.png"),
       "Fluxo do ADD estado a estado: quem fala e quem ouve no barramento W.", 15.0)
p("A figura traduz a tabela acima em transferências pelo barramento: em cada estado, "
  "uma fonte (à esquerda) escreve no barramento W e um destino (à direita) o amostra. "
  "Vê-se que a busca (T1–T3, em azul) é idêntica à de qualquer instrução e que a "
  "identidade do ADD só aparece na execução (T4–T6, em verde), culminando em A ← A + B.")
rich([("Observações. ", True),
      ("Em T1–T3 (busca) o processador nem sabe qual é a instrução — só em T3 o IR é "
       "carregado e o opcode passa a valer. Em T4 o operando (que é um ", False),
      ("endereço", True),
      (") é levado ao MAR; em T5 o dado lido vai para B; em T6 a ULA soma A + B e o "
       "resultado volta para A. Como A só é escrito no posedge de T6, seu novo valor "
       "aparece \"assentado\" já no T1 da próxima instrução — exatamente o efeito de "
       "\"um passo depois\" visível nas ondas. ", False)])
p("As outras instruções são variações deste esqueleto: LDA substitui o T5 por A ← RAM "
  "(e não usa T6); SUB é igual ao ADD com Su = 1 em T6; OUT resolve em T4 (saída ← A) e "
  "deixa T5/T6 vazios; HLT apenas trava em T4.")

# =====================================================================
# 9. CAMADA DE PLACA (fpga/)
# =====================================================================
doc.add_heading("9. Camada de placa (fpga/)", level=1)
p("A pasta fpga adapta o núcleo à DE10-Lite sem tocar no rtl. O top de placa "
  "(sap1_fpga.v) reúne quatro cuidados de integração.")

doc.add_heading("9.1. Geração e seleção de clock", level=2)
p("O clock de 50 MHz é rápido demais para observar a execução. Um divisor gera ~1 Hz "
  "(um estado T por segundo), e um seletor permite alternar para clock manual passo a "
  "passo (um pulso por toque de botão), ideal para depurar e gravar vídeo:")
code([
    "clock_divider #(.DIV(DIV)) u_div (.clk_in(clk50), .rst_n(rst_n), .clk_out(slow_clk));",
    "debouncer     #(.N(500_000)) u_db (.clk(clk50), .rst_n(rst_n),",
    "                                   .noisy(~KEY[1]), .clean(step_clk));",
    "wire cpu_clk = SW[0] ? step_clk : slow_clk;  // SW[0]: 0=auto(1Hz) 1=manual",
], caption="rtl/../fpga/sap1_fpga.v — clocks")
p("O divisor conta DIV−1 e alterna a saída (freq = 50 MHz / (2·DIV)). O debouncer filtra "
  "o botão com um sincronizador de dois estágios seguido de um contador: só aceita o "
  "novo nível após ele permanecer estável por N ciclos (~10 ms), evitando múltiplos "
  "pulsos por toque.")

doc.add_heading("9.2. Sincronizador de reset", level=2)
p("O botão de reset é assíncrono ao clock. Para evitar violações de recovery/removal, o "
  "reset é aplicado de forma assíncrona (assert imediato) mas liberado de forma síncrona "
  "(desassert passa por dois flip-flops). Isso, somado ao estado IDLE do núcleo (seção "
  "5.6), torna a partida totalmente determinística:")
code([
    "always @(posedge clk50 or negedge ext_rst_n) begin",
    "    if (!ext_rst_n) begin rs0 <= 0; rs1 <= 0; end  // assert assincrono",
    "    else            begin rs0 <= 1; rs1 <= rs0; end // desassert sincronizado",
    "end",
    "wire rst_n = rs1;   // ~CLR limpo para o nucleo",
], caption="fpga/sap1_fpga.v — sincronizador de reset")

doc.add_heading("9.3. Mapeamento dos displays", level=2)
p("Os seis displays de 7 segmentos mostram o funcionamento em tempo real. O opcode é "
  "traduzido para uma letra (em vez do hex cru) por um decodificador dedicado:")
table(
    ["Display", "Mostra", "Fonte"],
    [
        ["HEX5", "estado T atual (1–6; 0 = idle)", "tstate → número"],
        ["HEX4", "instrução como letra: L A 5 o H", "seg7_instr(opcode)"],
        ["HEX3–HEX2", "acumulador A (hex)", "acc (registrador real)"],
        ["HEX1–HEX0", "resultado (registrador de saída, hex)", "out_port"],
        ["LEDR[9]", "HLT (aceso ao parar)", "hlt"],
    ],
    widths=[2.8, 8.0, 4.6], fs=9.5,
)
code([
    "always @(*) case (opcode)",
    "    4'b0000: seg = 7'b1000111; // L  (LDA)",
    "    4'b0001: seg = 7'b0001000; // A  (ADD)",
    "    4'b0010: seg = 7'b0010010; // 5  (SUB, ~ 'S')",
    "    4'b1110: seg = 7'b0100011; // o  (OUT)",
    "    4'b1111: seg = 7'b0001001; // H  (HLT)",
    "    default: seg = 7'b1111111; // apagado",
    "endcase",
], caption="fpga/seg7_instr.v — opcode → letra (segmentos ativos em baixo)")
p("Há ainda um modo de depuração (SW[1] = 1) que redireciona os displays para mostrar o "
  "registrador B e o próprio barramento W a cada estado — útil para acompanhar as "
  "transferências internas na placa.", italic=True)

# =====================================================================
# 10. CÓDIGO-FONTE COMPLETO (MÓDULOS-CHAVE)
# =====================================================================
doc.add_heading("10. Código-fonte comentado dos módulos-chave", level=1)
p("Antes do código, esta seção fixa os conceitos de Verilog necessários para lê-lo; em "
  "seguida, cada módulo é apresentado com o código e uma explicação das construções "
  "usadas.")

doc.add_heading("10.1. Conceitos de Verilog usados", level=2)
p("Verilog descreve hardware, não um programa que executa passo a passo. Dois estilos "
  "convivem: lógica combinacional (a saída segue as entradas continuamente, como um "
  "circuito de portas) e lógica sequencial (o valor só muda na borda do clock, formando "
  "registradores). As construções abaixo são as usadas neste projeto:")
table(
    ["Construção", "O que significa"],
    [
        ["module … (portas)", "define um bloco de hardware e suas conexões externas"],
        ["input / output", "direção de cada porta do módulo"],
        ["wire", "fio: valor contínuo, dirigido por assign ou por outro bloco; não guarda estado"],
        ["reg", "sinal atribuído dentro de um bloco always; vira registrador (se sequencial) ou lógica (se combinacional)"],
        ["assign x = …", "atribuição contínua ⇒ lógica combinacional (o fio acompanha a expressão o tempo todo)"],
        ["always @(*)", "bloco combinacional: reavalia quando qualquer entrada muda"],
        ["always @(posedge clk)", "bloco sequencial: age na borda de subida do clock ⇒ gera flip-flops"],
        ["… or negedge clr_bar", "acrescenta o reset à lista de sensibilidade ⇒ reset assíncrono"],
        ["<=  (não-bloqueante)", "usado na lógica sequencial; as atribuições do bloco valem \"todas juntas\" na borda"],
        ["=  (bloqueante)", "usado na lógica combinacional (always @(*))"],
        ["localparam", "constante nomeada (ex.: opcodes, códigos de estado)"],
        ["case", "seleção múltipla — usada para decodificar opcode/estado"],
        ["{a, b}", "concatenação de bits (ex.: estender 4→8 bits)"],
        ["x[7:4]", "fatiamento: seleciona um intervalo de bits"],
        ["cond ? a : b", "operador ternário — um multiplexador de duas vias"],
    ],
    widths=[4.4, 11.0], fs=9.2,
)
lead("Reset assíncrono:", "quando clr_bar aparece na lista de sensibilidade (negedge "
     "clr_bar), o registrador zera no instante em que o reset é ativado, sem esperar o "
     "clock — por isso o if (!clr_bar) vem sempre primeiro (prioridade).")
lead("Por que <= (não-bloqueante):", "na lógica sequencial, todas as atribuições de um "
     "bloco devem enxergar os valores antigos e mudar simultaneamente na borda. O <= "
     "garante isso e evita corridas; o = (bloqueante) fica reservado à lógica "
     "combinacional.")
lead("reg não é \"registrador\" por si só:", "reg apenas indica que o sinal é atribuído "
     "dentro de um always. Se o always for @(posedge clk), a síntese cria um flip-flop; "
     "se for @(*), cria lógica combinacional.")

doc.add_heading("10.2. Program Counter", level=2)
code([
    "module program_counter (",
    "    input  wire       clk,",
    "    input  wire       clr_bar,   // ~CLR",
    "    input  wire       cp,        // Cp: habilita contagem",
    "    output reg  [3:0] pc_out",
    ");",
    "    always @(posedge clk or negedge clr_bar) begin",
    "        if (!clr_bar)   pc_out <= 4'b0000;",
    "        else if (cp)    pc_out <= pc_out + 4'b0001;",
    "    end",
    "endmodule",
], caption="rtl/program_counter.v")
p("Leitura do código: pc_out é output reg [3:0] — quatro bits (0 a 15) atribuídos dentro "
  "do always, logo é um registrador. O bloco reage a duas bordas: a subida do clock e a "
  "descida de clr_bar. A ordem das condições define a prioridade: primeiro if (!clr_bar) "
  "zera (reset assíncrono); senão, na borda de clock, se cp = 1, soma 1 (contagem); se "
  "cp = 0, não há atribuição e o valor é mantido — a síntese infere um flip-flop com "
  "enable. Ao passar de 15, os 4 bits voltam a 0 naturalmente (wraparound).")

doc.add_heading("10.3. Registrador de Instrução (decodifica opcode | operando)", level=2)
code([
    "module instruction_register (",
    "    input  wire clk, clr_bar, li_bar,",
    "    input  wire [7:0] bus_in,",
    "    output wire [3:0] opcode,    // nibble alto",
    "    output wire [3:0] operand    // nibble baixo",
    ");",
    "    reg [7:0] ir;",
    "    always @(posedge clk or negedge clr_bar) begin",
    "        if (!clr_bar)     ir <= 8'b0;",
    "        else if (!li_bar) ir <= bus_in;",
    "    end",
    "    assign opcode  = ir[7:4];",
    "    assign operand = ir[3:0];",
    "endmodule",
], caption="rtl/instruction_register.v")
p("Leitura do código: o registrador interno ir (8 bits) segue o mesmo padrão do PC — "
  "carrega bus_in quando li_bar = 0, senão mantém. A parte nova são os dois assign: como "
  "são atribuições contínuas, opcode e operand são wire e refletem sempre o conteúdo de "
  "ir. O fatiamento ir[7:4] pega os quatro bits altos (a instrução) e ir[3:0] os quatro "
  "baixos (o operando/endereço). Ou seja, a \"decodificação\" em dois campos não custa "
  "nenhum circuito extra — é apenas fiação para as duas metades do registrador.")
p("MAR, registrador B e registrador de saída seguem exatamente o padrão do acumulador "
  "(seção 5.2), mudando apenas a largura (MAR tem 4 bits) e o nome da carga.", italic=True)

doc.add_heading("10.4. Integração no topo (instanciação)", level=2)
code([
    "wire clk_cpu = clk;   // clock unico, SEM gating (HLT usa clock-enable)",
    "",
    "accumulator u_acc (.clk(clk_cpu), .clr_bar(clr_bar), .la_bar(la_bar),",
    "                   .bus_in(w_bus), .acc_out(acc_out));",
    "adder_subtractor u_alu (.su(su), .acc(acc_out), .breg(b_out), .result(alu_out));",
    "controller_sequencer u_ctrl (.clk(clk_cpu), .clr_bar(clr_bar), .opcode(ir_opcode),",
    "    .cp(cp), .ep(ep), .lm_bar(lm_bar), .ce_bar(ce_bar), .li_bar(li_bar),",
    "    .ei_bar(ei_bar), .la_bar(la_bar), .ea(ea), .su(su), .eu(eu),",
    "    .lb_bar(lb_bar), .lo_bar(lo_bar), .con(con), .ring(ring), .hlt(hlt));",
], caption="rtl/sap1_top.v — trecho de instanciação")
p("Leitura do código: instanciar um módulo é criar uma cópia física do bloco e ligar "
  "seus pinos. A forma nome_do_modulo nome_da_instancia (.porta(sinal), …) faz a conexão "
  "por nome — por exemplo, .bus_in(w_bus) liga a entrada bus_in do acumulador ao "
  "barramento. O mesmo w_bus chega a vários blocos ao mesmo tempo (todos podem ler), mas "
  "só um escreve por vez, graças ao mux. Note que clk_cpu é apenas o clock de entrada "
  "sem nenhuma porta lógica no caminho — coerente com a decisão de não fazer gating (o "
  "HLT congela via clock-enable, não desligando o clock).")

# =====================================================================
# 11. EXECUÇÃO COMPLETA DE UM PROGRAMA (PASSO A PASSO)
# =====================================================================
doc.add_heading("11. Execução completa de um programa (passo a passo)", level=1)
p("Para fechar o funcionamento, esta seção acompanha um programa inteiro do início ao "
  "fim, amarrando o que foi visto em separado — estados, palavra de controle e datapath. "
  "Usamos o programa da multiplicação 3 × 4 = 12 (somas repetidas), por ser curto e "
  "visualmente claro. A memória contém, nos dados, mem[11] = 3 e mem[12..15] = 0.")
p("A tabela mostra cada instrução com a ação executada (resumida pelos estados usados) e "
  "os valores dos registradores após terminá-la. Cada instrução consome os seis estados "
  "T1–T6; a busca (T1–T3) é sempre a mesma e foi omitida da coluna de ação por brevidade.")
table(
    ["Endereço", "Instrução", "Ação (execução)", "A após", "Saída"],
    [
        ["0", "LDA 11", "T4 MAR←11; T5 A←mem[11]=3", "3", "0"],
        ["1", "ADD 11", "T4 MAR←11; T5 B←3; T6 A←A+B", "6", "0"],
        ["2", "ADD 11", "T5 B←3; T6 A←A+B", "9", "0"],
        ["3", "ADD 11", "T5 B←3; T6 A←A+B", "12", "0"],
        ["4", "OUT", "T4 saída←A", "12", "12"],
        ["5", "SUB 12", "T5 B←mem[12]=0; T6 A←A−B", "12", "12"],
        ["6", "ADD 13", "T5 B←0; T6 A←A+B", "12", "12"],
        ["7", "SUB 14", "T5 B←0; T6 A←A−B", "12", "12"],
        ["8", "ADD 15", "T5 B←0; T6 A←A+B", "12", "12"],
        ["9", "OUT", "T4 saída←A", "12", "12"],
        ["10", "HLT", "T4 trava o anel (halted←1)", "12", "12"],
    ],
    widths=[1.9, 2.2, 6.4, 1.7, 1.6], fs=9.3,
)
p("Lendo a tabela: as quatro primeiras instruções constroem a multiplicação somando o "
  "3 repetidamente (A: 3 → 6 → 9 → 12) — é assim que o SAP-1 \"multiplica\", já que não "
  "há instrução de multiplicação nem laços. Em seguida, o OUT (endereço 4) copia 12 para "
  "o visor. As instruções 5 a 8 operam com dados iguais a zero, de modo que A não muda — "
  "servem apenas para preencher o programa até o segundo OUT, que confirma o resultado. "
  "Por fim, o HLT congela a máquina em T4, mantendo 12 no visor. Esse é exatamente o "
  "comportamento observado nas formas de onda da Parte II.")

doc.add_heading("11.1. As demais instruções, em resumo", level=2)
p("O ADD e o SUB já foram rastreados ciclo a ciclo na seção 8. As outras três, que usam "
  "menos estados, completam o quadro:")
lead("LDA end — carrega da memória.",
     "Após a busca, no T4 o operando vai ao MAR e no T5 o dado lido da RAM é escrito no "
     "acumulador (A ← mem[end]). Termina em T5 (o T6 fica vazio). É a única forma de "
     "trazer um valor da memória para dentro do processador.")
lead("OUT — exibe o acumulador.",
     "Após a busca, resolve-se tudo no T4: o acumulador escreve no barramento (Ea) e o "
     "registrador de saída o captura (~Lo). Não altera A nem a memória — apenas mostra o "
     "valor no visor. T5 e T6 ficam vazios.")
lead("HLT — para o processador.",
     "Após a busca, no T4 o flag de parada é travado (halted ← 1) e o contador de anel "
     "deixa de avançar (congela em T4). Todos os registradores mantêm seus valores e o "
     "clock continua correndo; só um reset reinicia a execução.")

# =====================================================================
# 12. CONCLUSÃO
# =====================================================================
doc.add_heading("12. Conclusão", level=1)
p("A implementação seguiu três eixos. Modularidade: um arquivo por bloco, com o topo "
  "apenas interligando-os pelo barramento, e uma separação estrita rtl × fpga. Boas "
  "práticas de FPGA no nível de hardware: barramento por multiplexador (sem tri-state), "
  "reset assíncrono com sincronizador de desassert, ausência total de gating de clock "
  "(HLT por clock-enable) e disciplina de duas bordas para eliminar corrida entre "
  "controle e dados. Clareza didática: palavra de controle transcrita da tabela de "
  "microoperações, máquina de estados explícita e disponível em duas formas equivalentes "
  "(anel one-hot e FSM clássica). O resultado é um núcleo sintetizável, determinístico na "
  "partida e legível — cuja unidade de controle, o ponto mais delicado, está descrita de "
  "forma direta e verificável. A validação por formas de onda consta do relatório de "
  "verificação por simulação.")

doc.save(OUT)
print("gerado:", OUT)
